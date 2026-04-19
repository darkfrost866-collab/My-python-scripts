from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from pathlib import Path
import os

try:
    from docxtpl import DocxTemplate
    TEMPLATE_AVAILABLE = True
except:
    TEMPLATE_AVAILABLE = False

def create_resume_docx(job, tailored_bullets, output_dir):
    filename = output_dir / f"Raouf_Mayahi_Resume_{job['company'].replace(' ','_').replace('/','-')}.docx"
    
    # Use template if available
    template_path = Path('resume_template.docx')
    if TEMPLATE_AVAILABLE and template_path.exists():
        try:
            doc = DocxTemplate(str(template_path))
            context = {
                'company': job['company'],
                'job_title': job['title'],
                'date': '2026',
            }
            doc.render(context)
            doc.save(str(filename))
            return filename
        except:
            pass
    
    # Fallback to professional format
    doc = Document()
    
    # Header
    header = doc.add_heading('RAOUF MAYAHI', 0)
    header.alignment = 1
    
    contact = doc.add_paragraph('Production Supervisor | Operations Leader\nYork, ON | (416) 897-7889 | raoufmayahi@gmail.com\nlinkedin.com/in/raoufmayahi')
    contact.alignment = 1
    
    doc.add_paragraph()
    
    # Target
    target = doc.add_heading(f'TARGET POSITION: {job["title"].upper()}', 2)
    
    # Summary
    doc.add_heading('PROFESSIONAL SUMMARY', 1)
    summary = doc.add_paragraph(
        'Results-driven Operations Leader with 14+ years of progressive experience in heavy manufacturing, '
        'custom fabrication, and high-precision assembly environments. Demonstrated expertise in Lean Six Sigma implementation, '
        'root cause analysis, and HSE leadership within unionized workforces. Proven ability to drive operational excellence '
        'while maintaining uncompromising safety standards. Certified TSSA (1G-6G) and CWB Welding Inspector with comprehensive '
        'knowledge of ISO 9001, IATF 16949, and AS9100 quality management systems.'
    )
    summary.alignment = 3
    
    # Core Competencies
    doc.add_heading('CORE COMPETENCIES', 1)
    competencies = [
        'Lean Six Sigma & Continuous Improvement | Root Cause Analysis | 5S Implementation',
        'HSE Leadership & OHSA Compliance | JHSC Certified | Zero LTI Achievement',
        'Union Labor Relations | Dispute Resolution | Team Leadership (45+ personnel)',
        'TSSA Certified (1G-6G) | CWB Welding Inspector | Blueprint Reading',
        'ISO 9001/IATF 16949/AS9100 | Quality Management Systems | Process Auditing',
        'Production Planning | Material Procurement | Budget Optimization | KPI Management'
    ]
    for comp in competencies:
        doc.add_paragraph(comp, style='List Bullet')
    
    # Key Achievements
    doc.add_heading('KEY ACHIEVEMENTS', 1)
    for bullet in tailored_bullets.split('\n'):
        if bullet.strip():
            clean_bullet = bullet.replace('•', '').strip()
            doc.add_paragraph(clean_bullet, style='List Bullet')
    
    # Experience
    doc.add_heading('PROFESSIONAL EXPERIENCE', 1)
    
    doc.add_heading('ACTIVE DYNAMICS | Production Supervisor | 2024 – Present', 2)
    exp1 = [
        'Provide strategic leadership for 45+ unionized personnel across multi-shift manufacturing operations',
        'Engineered 15% increase in shift effectiveness through Lean Six Sigma implementation',
        'Achieved zero Lost Time Injuries over 24 months via rigorous HSE program management',
        'Resolved 90% of labor disputes through collaborative engagement and proactive communication',
        'Optimized material procurement processes, reducing waste by 12%'
    ]
    for item in exp1:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('HOWDEN CANADA | Lead Hand | 2018 – 2022', 2)
    exp2 = [
        'Ensured 100% adherence to ISO quality management systems and regulatory requirements',
        'Supported 12% facility capacity expansion through process optimization initiatives',
        'Mentored junior staff in welding best practices and safety protocols',
        'Conducted root cause analysis for production variances, implementing corrective actions'
    ]
    for item in exp2:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('BDI CANADA | Welder/Fabricator | 2012 – 2018', 2)
    exp3 = [
        'Performed precision welding and fabrication in accordance with TSSA and CWB standards',
        'Interpreted complex blueprints and technical drawings for custom manufacturing projects',
        'Maintained exemplary safety record in heavy industrial environment'
    ]
    for item in exp3:
        doc.add_paragraph(item, style='List Bullet')
    
    # Education & Certifications
    doc.add_heading('EDUCATION & CERTIFICATIONS', 1)
    certs = [
        'Engineering Technology Diploma | Humber College',
        'TSSA Certified Welder (1G, 2G, 3G, 4G, 5G, 6G)',
        'CWB Level 2 Welding Inspector',
        'Lean Six Sigma Green Belt',
        'JHSC Certification | OHSA Training',
        'WHMIS | Working at Heights | Forklift Certified'
    ]
    for cert in certs:
        doc.add_paragraph(cert, style='List Bullet')
    
    doc.save(str(filename))
    return filename

def create_resume_pdf(job, tailored_bullets, output_dir):
    filename = output_dir / f"Raouf_Mayahi_Resume_{job['company'].replace(' ','_')}.pdf"
    doc = SimpleDocTemplate(str(filename), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    story.append(Paragraph("<b>RAOUF MAYAHI</b>", styles['Title']))
    story.append(Paragraph("Production Supervisor | (416) 897-7889 | raoufmayahi@gmail.com", styles['Normal']))
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"<b>{job['title']} - {job['company']}</b>", styles['Heading2']))
    story.append(Spacer(1, 12))
    
    for bullet in tailored_bullets.split('\n'):
        if bullet.strip():
            story.append(Paragraph(bullet, styles['Normal']))
    
    doc.build(story)
    return filename

def create_cover_pdf(job, cover_text, output_dir):
    filename = output_dir / f"Raouf_Mayahi_Cover_{job['company'].replace(' ','_')}.pdf"
    doc = SimpleDocTemplate(str(filename), pagesize=letter, leftMargin=72, rightMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    story = []
    
    for line in cover_text.split('\n'):
        if line.strip():
            story.append(Paragraph(line.replace(' ', '&nbsp;' * 0), styles['Normal']))
        else:
            story.append(Spacer(1, 12))
    
    doc.build(story)
    return filename

def create_cover_docx(job, cover_text, output_dir):
    filename = output_dir / f"Raouf_Mayahi_Cover_{job['company'].replace(' ','_')}.docx"
    doc = Document()
    
    for line in cover_text.split('\n'):
        doc.add_paragraph(line)
    
    doc.save(str(filename))
    return filename